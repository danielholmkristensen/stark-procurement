#!/usr/bin/env python3
"""
STARK Procurement Scope Document — Word Generator
Creates an on-brand document in Agentic Agency style (Off-white/Cement palette)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Output path
OUTPUT_PATH = "/Users/dhk/Projects/STARK_Procurement/app/docs/STARK_Procurement_Scope_Document_Kim_Christensen.docx"

# Agentic Agency Brand Colors - LIGHT THEME
AA_BLACK = RGBColor(0x00, 0x00, 0x00)
AA_CEMENT = RGBColor(0xE6, 0xE6, 0xE1)  # Off-white/cream - primary background
AA_CEMENT_DARK = RGBColor(0xD0, 0xD0, 0xCB)  # Slightly darker cement for contrast
AA_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
AA_MID_GRAY = RGBColor(0x66, 0x66, 0x66)
AA_LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = AA_MID_GRAY

def create_document():
    """Create the full Word document"""
    doc = Document()

    # Set up default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = AA_BLACK

    # ═══════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════

    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("STARK PROCUREMENT")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = AA_BLACK

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scope Definition Document")
    run.font.size = Pt(24)
    run.font.color.rgb = AA_DARK_GRAY

    # Recipient
    doc.add_paragraph()
    doc.add_paragraph()
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = recipient.add_run("Prepared for Kim Christensen")
    run.font.size = Pt(16)
    run.font.color.rgb = AA_MID_GRAY

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = AA_MID_GRAY

    # Vendor attribution
    for _ in range(8):
        doc.add_paragraph()

    vendor = doc.add_paragraph()
    vendor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = vendor.add_run("Delivered by Agentic Agency")
    run.font.size = Pt(12)
    run.font.color.rgb = AA_MID_GRAY

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Engineering > Prompting")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = AA_MID_GRAY

    # Page break
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # INTRODUCTION: OUR APPROACH
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Our Approach: Engineering Excellence at Speed", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    intro = doc.add_paragraph()
    intro.add_run(
        "At Agentic Agency, we believe in a fundamental truth: great software comes from great engineering, not from shortcuts. "
        "We've pioneered a delivery model that combines the speed of modern AI-assisted development with the rigor and quality "
        "of traditional software craftsmanship. The result is exceptional software delivered in compressed timelines—without compromise."
    )
    intro.paragraph_format.space_after = Pt(12)

    # Values section
    h2 = doc.add_heading("Our Values", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    values = [
        ("Quality Without Compromise", "We write code that you would be proud to show your best engineers. Clean, maintainable, well-tested, and documented."),
        ("Speed Through Intelligence", "We leverage AI as a force multiplier for skilled engineers—not as a replacement for engineering judgment."),
        ("Transparency in Execution", "You see everything we build, as we build it. No black boxes. No surprises."),
        ("Partnership Over Transaction", "We succeed when you succeed. Our goal is a long-term relationship, not a one-off sale."),
    ]

    for title, desc in values:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    add_horizontal_line(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # UNDERSTANDING THE LANDSCAPE: DELIVERY APPROACHES
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Understanding the Landscape: Software Delivery Approaches", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    landscape_intro = doc.add_paragraph()
    landscape_intro.add_run(
        "The software development industry offers various approaches to building systems. "
        "Understanding these approaches—their strengths and limitations—is essential for making informed decisions."
    )
    landscape_intro.paragraph_format.space_after = Pt(16)

    # LOW CODE
    h2 = doc.add_heading("Low-Code Platforms", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    lowcode = doc.add_paragraph()
    lowcode.add_run("Definition: ").bold = True
    lowcode.add_run(
        "Low-code platforms (e.g., OutSystems, Mendix, Power Platform) provide visual, drag-and-drop interfaces "
        "to build applications. They abstract away underlying code, allowing 'citizen developers' to create software without traditional programming."
    )
    lowcode.paragraph_format.space_after = Pt(8)

    lowcode_pros = doc.add_paragraph()
    lowcode_pros.add_run("Strengths: ").bold = True
    lowcode_pros.add_run(
        "Rapid prototyping, reduced need for specialized developers, accessible to non-technical users, "
        "good for standardized CRUD applications."
    )
    lowcode_pros.paragraph_format.space_after = Pt(8)

    lowcode_cons = doc.add_paragraph()
    lowcode_cons.add_run("Limitations: ").bold = True
    lowcode_cons.add_run(
        "Vendor lock-in, limited customization, performance constraints at scale, hidden complexity that surfaces later, "
        "ongoing licensing costs, difficulty integrating with complex enterprise systems, and often produces code that is "
        "unmaintainable outside the platform."
    )
    lowcode_cons.paragraph_format.space_after = Pt(16)

    # VIBECODING
    h2 = doc.add_heading("Vibecoding (Prompt-Driven Development)", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    vibe = doc.add_paragraph()
    vibe.add_run("Definition: ").bold = True
    vibe.add_run(
        "Vibecoding refers to the practice of using conversational AI tools (like ChatGPT, Claude, or Cursor) "
        "to generate code through prompts, without deep understanding of the underlying technology. "
        "The developer 'vibes' with the AI, accepting generated code with minimal review."
    )
    vibe.paragraph_format.space_after = Pt(8)

    vibe_pros = doc.add_paragraph()
    vibe_pros.add_run("Strengths: ").bold = True
    vibe_pros.add_run(
        "Extremely fast initial output, low barrier to entry, can produce working prototypes quickly, "
        "democratizes access to code generation."
    )
    vibe_pros.paragraph_format.space_after = Pt(8)

    vibe_cons = doc.add_paragraph()
    vibe_cons.add_run("Limitations: ").bold = True
    vibe_cons.add_run(
        "Inconsistent quality, security vulnerabilities, lack of architectural coherence, technical debt accumulation, "
        "difficulty maintaining generated code, hallucinated solutions that don't actually work, no understanding of "
        "the 'why' behind decisions, and brittleness under real-world conditions. Vibecoding produces artifacts, not systems."
    )
    vibe_cons.paragraph_format.space_after = Pt(16)

    # TRADITIONAL DEVELOPMENT
    h2 = doc.add_heading("Traditional Software Development", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    trad = doc.add_paragraph()
    trad.add_run("Definition: ").bold = True
    trad.add_run(
        "Conventional software engineering with skilled developers, established methodologies (Agile, Scrum, Waterfall), "
        "code reviews, testing practices, and deliberate architectural decisions. Code is written by humans who understand it."
    )
    trad.paragraph_format.space_after = Pt(8)

    trad_pros = doc.add_paragraph()
    trad_pros.add_run("Strengths: ").bold = True
    trad_pros.add_run(
        "High quality when done well, full control, deep understanding of the codebase, maintainability, "
        "appropriate for complex systems, proven methodologies, transferable skills."
    )
    trad_pros.paragraph_format.space_after = Pt(8)

    trad_cons = doc.add_paragraph()
    trad_cons.add_run("Limitations: ").bold = True
    trad_cons.add_run(
        "Time-intensive, expensive, dependent on finding skilled talent, can be slow to adapt, "
        "often burdened by process overhead, communication gaps between business and technical teams."
    )
    trad_cons.paragraph_format.space_after = Pt(16)

    add_horizontal_line(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # AGENTIC ENGINEERING: OUR APPROACH
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Agentic Engineering: A Different Model", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    ae_intro = doc.add_paragraph()
    ae_intro.add_run(
        "Agentic Engineering is neither low-code, nor vibecoding, nor traditional development. "
        "It represents a new paradigm where AI augments—rather than replaces—expert human judgment."
    )
    ae_intro.paragraph_format.space_after = Pt(12)

    ae_core = doc.add_paragraph()
    ae_core.add_run("The Core Principle: ").bold = True
    ae_core.add_run(
        "Skilled engineers remain in the loop at every decision point. AI handles the mechanical aspects of code generation, "
        "pattern matching, and documentation—freeing engineers to focus on architecture, edge cases, security, and the "
        "nuanced decisions that determine software quality."
    )
    ae_core.paragraph_format.space_after = Pt(16)

    h2 = doc.add_heading("How Agentic Engineering Differs", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    # Comparison table
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Aspect", "Low-Code", "Vibecoding", "Traditional", "Agentic Engineering"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "E6E6E1")
        cell.paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    data = [
        ["Speed", "Fast", "Very Fast", "Slow", "Fast"],
        ["Quality", "Medium", "Low-Variable", "High (if done well)", "High (guaranteed)"],
        ["Customization", "Limited", "Unlimited (chaotic)", "Full control", "Full control"],
        ["Maintainability", "Platform-dependent", "Poor", "High", "High"],
        ["Cost", "Medium + licensing", "Low (initial)", "High", "Competitive"],
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_data
            if col_idx == 4:  # Highlight our column
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    h2 = doc.add_heading("The ADAPT Methodology", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    adapt_intro = doc.add_paragraph()
    adapt_intro.add_run("ADAPT: ").bold = True
    adapt_intro.add_run("Agentic Development with Artifact Persistence & Testing")
    adapt_intro.paragraph_format.space_after = Pt(12)

    adapt_desc = doc.add_paragraph()
    adapt_desc.add_run(
        "ADAPT is a methodology for AI-assisted software delivery that ensures knowledge compounds across sessions, "
        "quality is enforced through test-gated tasks, and work is structured for maximum parallelization."
    )
    adapt_desc.paragraph_format.space_after = Pt(12)

    adapt_items = [
        ("Agentic", "AI agents work autonomously on well-defined tasks within clear boundaries. Human engineers set direction, review output, and make architectural decisions."),
        ("Development", "Real engineering—not prototyping. Production-grade code with proper error handling, type safety, and maintainability."),
        ("Artifact", "Every decision, lesson, and context is persisted in structured knowledge stores. Nothing is lost between sessions."),
        ("Persistence", "Session continuity through shared context logs, lessons learned, and development diaries. The next session picks up exactly where this one left off."),
        ("Testing", "Test-gated task completion. No task is marked done until tests pass. Quality is enforced, not hoped for."),
    ]

    for term, desc in adapt_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.25)

    add_horizontal_line(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PRICING CONTEXT
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("An Invitation to Experience the Model", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    pricing_intro = doc.add_paragraph()
    pricing_intro.add_run(
        "The price point for this engagement reflects a strategic decision, not a market rate. "
        "We are extending an invitation to experience Agentic Engineering firsthand."
    )
    pricing_intro.paragraph_format.space_after = Pt(12)

    pricing_why = doc.add_paragraph()
    pricing_why.add_run("Why this price? ").bold = True
    pricing_why.add_run(
        "Agentic Engineering is a relatively unknown delivery model. Many organizations are skeptical—and rightfully so, "
        "given the hype and disappointment cycles in technology. We believe the best way to demonstrate the model is to "
        "deliver results that speak for themselves."
    )
    pricing_why.paragraph_format.space_after = Pt(12)

    pricing_not = doc.add_paragraph()
    pricing_not.add_run("What this is NOT: ").bold = True
    pricing_not.add_run(
        "This is not a race to the bottom. We are not undercutting the market to win volume. This price is a one-time "
        "showcase opportunity—an investment in demonstrating capability and building a reference relationship. "
        "Future engagements will reflect the true value delivered."
    )
    pricing_not.paragraph_format.space_after = Pt(12)

    pricing_get = doc.add_paragraph()
    pricing_get.add_run("What you get: ").bold = True
    pricing_get.add_run(
        "The same quality, rigor, and professionalism as our full-rate engagements. No shortcuts. No reduced scope. "
        "No junior resources. You receive our best work."
    )

    add_horizontal_line(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # OUR GUARANTEES
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Our Guarantees", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    guarantee_intro = doc.add_paragraph()
    guarantee_intro.add_run(
        "We don't ask you to trust our judgment of quality. We ask you to define it."
    )
    guarantee_intro.paragraph_format.space_after = Pt(16)

    # Code Quality
    h2 = doc.add_heading("Code Quality: Show Us Your Standard", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    code_qual = doc.add_paragraph()
    code_qual.add_run(
        "Show us examples of what you consider great code—or even beautiful code. It might be from your existing codebase, "
        "from an open-source project you admire, or from a book you reference. We will study it, understand its patterns, "
        "and adhere to that standard throughout our delivery."
    )
    code_qual.paragraph_format.space_after = Pt(8)

    code_examples = doc.add_paragraph()
    code_examples.add_run("Examples might include: ").bold = True
    code_examples.add_run(
        "A Rails project with clean service objects, a React codebase with excellent component composition, "
        "a Go service with idiomatic error handling, or documentation-driven development patterns you value."
    )
    code_examples.paragraph_format.space_after = Pt(16)

    # Documentation Quality
    h2 = doc.add_heading("Documentation: Show Us What Impressed You", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    doc_qual = doc.add_paragraph()
    doc_qual.add_run(
        "Show us examples of solution architecture documents or technical documentation that made you say 'wow'—regardless "
        "of the context or industry. We will ADAPT to that standard. If you admire Stripe's API documentation, "
        "we'll structure ours similarly. If you prefer AWS's well-architected framework style, we'll match it."
    )
    doc_qual.paragraph_format.space_after = Pt(16)

    # Delivery Walkthrough
    h2 = doc.add_heading("Delivery Walkthrough: Full Transparency", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    walkthrough = doc.add_paragraph()
    walkthrough.add_run(
        "At the end of our delivery, we will conduct a full one-hour walkthrough with your technical team. "
        "We will demonstrate that we have strictly delivered within the agreed parameters—no more, no less. "
        "Every screen, every integration, every business rule will be shown working against real data."
    )
    walkthrough.paragraph_format.space_after = Pt(8)

    walkthrough_scope = doc.add_paragraph()
    walkthrough_scope.add_run(
        "This document defines those parameters. Everything marked 'IN SCOPE' will be delivered. "
        "Everything marked 'OUT OF SCOPE' will not be delivered—and we will not charge extra for scope we never agreed to."
    )

    add_horizontal_line(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PARTNERSHIP VISION
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Our Aim: Partnership", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    partner_intro = doc.add_paragraph()
    partner_intro.add_run(
        "We are not looking for a transaction. We are looking for a partnership."
    )
    partner_intro.paragraph_format.space_after = Pt(12)

    partner_detail = doc.add_paragraph()
    partner_detail.add_run(
        "STARK Group's digital transformation journey spans multiple systems, teams, and years. "
        "We believe Agentic Engineering can be a valuable capability in that journey—not just for procurement, "
        "but across the portfolio. This engagement is an opportunity for both parties to evaluate fit."
    )
    partner_detail.paragraph_format.space_after = Pt(12)

    partner_success = doc.add_paragraph()
    partner_success.add_run("Success for us means: ").bold = True
    partner_success.add_run(
        "You complete this project with confidence in our capability, with a system you're proud of, "
        "and with the desire to continue working together."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SCOPE DEFINITION BEGINS
    # ═══════════════════════════════════════════════════════════════════════

    h1 = doc.add_heading("Scope Definition: STARK Procurement", level=1)
    h1.runs[0].font.color.rgb = AA_BLACK

    scope_intro = doc.add_paragraph()
    scope_intro.add_run("Document Purpose: ").bold = True
    scope_intro.add_run("Vendor-ready, MECE scope definition for fixed-price delivery")

    scope_version = doc.add_paragraph()
    scope_version.add_run("Version: ").bold = True
    scope_version.add_run("1.0 DRAFT | Status: For Discussion")

    add_horizontal_line(doc)

    # EXECUTIVE SUMMARY
    h2 = doc.add_heading("Executive Summary", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("What We Are Building", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    building = doc.add_paragraph()
    building.add_run(
        "A Procurement System that handles the complete Purchase-to-Pay workflow for STARK Group Denmark:"
    )
    building.paragraph_format.space_after = Pt(8)

    flow = doc.add_paragraph()
    flow.add_run("PR Ingestion → PO Generation → Supplier Communication → Invoice Matching → Approval")
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow.runs[0].bold = True
    flow.paragraph_format.space_after = Pt(16)

    h3 = doc.add_heading("What We Are NOT Building", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    not_building = [
        "Warehouse Management (NYCE team)",
        "Financial Accounting (SAP Finance team)",
        "E-commerce Platform (ECom team)",
        "Demand Planning (Relex team)",
        "Contract Management (Icertis team)",
    ]

    for item in not_building:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    # Scale table
    h3 = doc.add_heading("Scale", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    scale_table = doc.add_table(rows=6, cols=2)
    scale_table.style = 'Table Grid'

    scale_data = [
        ("Metric", "Annual Volume"),
        ("Purchase Orders", "750,000"),
        ("PO Value", "€1.1B"),
        ("Suppliers", "10,000"),
        ("Branches", "84"),
        ("System Users", "30 procurement staff"),
    ]

    for row_idx, (col1, col2) in enumerate(scale_data):
        scale_table.rows[row_idx].cells[0].text = col1
        scale_table.rows[row_idx].cells[1].text = col2
        if row_idx == 0:
            scale_table.rows[row_idx].cells[0].paragraphs[0].runs[0].bold = True
            scale_table.rows[row_idx].cells[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(scale_table.rows[row_idx].cells[0], "E6E6E1")
            set_cell_shading(scale_table.rows[row_idx].cells[1], "E6E6E1")
            scale_table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.color.rgb = AA_BLACK
            scale_table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()
    add_horizontal_line(doc)

    # USERS & ROLES
    h2 = doc.add_heading("Users & Roles", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("User Population", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    users_table = doc.add_table(rows=4, cols=3)
    users_table.style = 'Table Grid'

    users_data = [
        ("Role", "Count (est.)", "Primary Function"),
        ("Buyer", "~20", "PR review, PO creation, bundling, invoice matching"),
        ("Approver", "~8", "PO approval, invoice approval, exception handling"),
        ("Admin", "~2", "System configuration, user support, reporting"),
    ]

    for row_idx, row_data in enumerate(users_data):
        for col_idx, cell_data in enumerate(row_data):
            users_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                users_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(users_table.rows[row_idx].cells[col_idx], "E6E6E1")
                users_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()

    # Role definitions
    h3 = doc.add_heading("Role Definitions", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    roles = [
        ("Buyer", "Day-to-day procurement operations. Reviews PRs, creates/sends POs, matches invoices, resolves discrepancies within tolerance. Cannot approve above their personal authority limit."),
        ("Approver", "Decision authority for high-value transactions. Approves POs and invoices above threshold. Handles escalated discrepancies. May also perform Buyer functions."),
        ("Admin", "System configuration and support. Manages thresholds, supplier preferences, cut-off times, user authority limits, and delegations. No transactional authority beyond Approver role."),
    ]

    for role, desc in roles:
        p = doc.add_paragraph()
        run = p.add_run(f"{role}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    add_horizontal_line(doc)

    # DOMAIN STRUCTURE
    h2 = doc.add_heading("Scope Structure", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    scope_struct = doc.add_paragraph()
    scope_struct.add_run("This document defines scope across 5 mutually exclusive domains:")
    scope_struct.paragraph_format.space_after = Pt(12)

    domains = [
        ("Domain 1: User Interface", "Screens, components, interactions"),
        ("Domain 2: Business Logic", "Rules, calculations, workflows"),
        ("Domain 3: Data Layer", "Storage, queries, synchronization"),
        ("Domain 4: Integrations", "Inbound data, outbound data, APIs"),
        ("Domain 5: Operations", "Infrastructure, deployment, monitoring"),
    ]

    for domain, desc in domains:
        p = doc.add_paragraph()
        run = p.add_run(f"{domain}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    scope_method = doc.add_paragraph()
    scope_method.add_run("\nFor each domain, we define:")
    scope_method.paragraph_format.space_after = Pt(8)

    for item in ["IN — We build, we own, we deliver", "OUT — Someone else owns, we don't touch", "BOUNDARY — Exact handoff point"]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    add_horizontal_line(doc)

    # DOMAIN 1: USER INTERFACE
    h2 = doc.add_heading("Domain 1: User Interface", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("1.1 Screens — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    screens_table = doc.add_table(rows=17, cols=3)
    screens_table.style = 'Table Grid'

    screens_data = [
        ("Screen ID", "Name", "Purpose"),
        ("A1", "PR Inbox", "View and manage incoming purchase requests"),
        ("A2", "PR Detail", "Individual PR with line items"),
        ("B1", "PO List", "View and manage purchase orders"),
        ("B2", "PO Detail", "Individual PO with supplier, items, status"),
        ("B3", "PO Kanban", "Visual pipeline (Draft → Sent → Confirmed → Received)"),
        ("C1", "Bundling Workspace", "Group PRs into optimized POs"),
        ("D1", "Invoice List", "View and manage supplier invoices"),
        ("D2", "Invoice Detail", "Individual invoice with matching status"),
        ("D3", "Match Results", "Side-by-side PO vs Invoice comparison"),
        ("D4", "Discrepancy Queue", "Invoices requiring manual resolution"),
        ("E1", "Approval Queue", "Items awaiting approval"),
        ("E2", "Approval History", "Audit trail of decisions"),
        ("F1", "Supplier List", "Supplier master data"),
        ("F2", "Supplier Detail", "Individual supplier with performance"),
        ("G1", "Dashboard", "Morning briefing, action items, KPIs"),
        ("H1", "Settings", "User preferences, thresholds"),
    ]

    for row_idx, row_data in enumerate(screens_data):
        for col_idx, cell_data in enumerate(row_data):
            screens_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                screens_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(screens_table.rows[row_idx].cells[col_idx], "E6E6E1")
                screens_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    total_screens = doc.add_paragraph()
    total_screens.add_run("\nTotal: 16 screens")
    total_screens.runs[0].bold = True

    add_horizontal_line(doc)

    # DOMAIN 2: BUSINESS LOGIC
    h2 = doc.add_heading("Domain 2: Business Logic", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("2.1 PR Processing — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    pr_rules = [
        ("PR Validation", "Validate incoming PR data — Required fields present, valid supplier ID, positive quantities"),
        ("PR Deduplication", "Detect duplicate PRs — Same source + reference = reject with 409"),
        ("Source Routing", "Handle PR differently by source — Relex: batch at cut-off, ECom: immediate, SalesApp: per timing field"),
        ("Urgency Calculation", "Assign escalation level — Based on value, age, need-by date"),
    ]

    for rule, desc in pr_rules:
        p = doc.add_paragraph()
        run = p.add_run(f"{rule}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    h3 = doc.add_heading("2.2 PO Generation — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    po_rules = [
        ("Single PR → PO", "Convert one PR to one PO — 1:1 mapping, all line items transferred"),
        ("Bundling", "Combine PRs into one PO — Same supplier + same location + compatible timing"),
        ("Packet Specification", "Mark items per-PR in bundled PO — Supplier receives clear picking instructions"),
        ("Cut-off Timing", "Send PO at supplier cut-off — Configurable per supplier, default to immediate"),
    ]

    for rule, desc in po_rules:
        p = doc.add_paragraph()
        run = p.add_run(f"{rule}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    h3 = doc.add_heading("2.5 Invoice Matching — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    inv_rules = [
        ("2-Way Match", "Compare Invoice to PO — Match on PO number, line items, quantities, prices"),
        ("Tolerance Rules", "Allow minor variances — Qty: ±2%, Price: ±1%, Total: ±DKK 100"),
        ("Discrepancy Detection", "Flag mismatches — Qty mismatch, price mismatch, missing PO"),
        ("Auto-Approve", "Pass matching invoices — Full match within tolerance = auto-approve"),
    ]

    for rule, desc in inv_rules:
        p = doc.add_paragraph()
        run = p.add_run(f"{rule}: ")
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(6)

    add_horizontal_line(doc)

    # DOMAIN 4: INTEGRATIONS
    h2 = doc.add_heading("Domain 4: Integrations", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    int_arch = doc.add_paragraph()
    int_arch.add_run("Architecture Decision: ").bold = True
    int_arch.add_run(
        "Kafka-native integration with per-source/per-consumer topic modularity. "
        "REST fallback for sources that cannot use Kafka. Customer provisions topics in existing STARK Kafka cluster; "
        "vendor delivers specs and consumer/producer code."
    )
    int_arch.paragraph_format.space_after = Pt(12)

    h3 = doc.add_heading("4.2 Inbound: PR Sources — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    inbound_table = doc.add_table(rows=5, cols=3)
    inbound_table.style = 'Table Grid'

    inbound_data = [
        ("Source", "Topic", "Consumer"),
        ("Relex", "stark.procurement.inbound.prs.relex", "PRRelexConsumer"),
        ("ECom", "stark.procurement.inbound.prs.ecom", "PREComConsumer"),
        ("SalesApp", "stark.procurement.inbound.prs.salesapp", "PRSalesAppConsumer"),
        ("Aspect4", "stark.procurement.inbound.pos.aspect4", "POAspect4Consumer"),
    ]

    for row_idx, row_data in enumerate(inbound_data):
        for col_idx, cell_data in enumerate(row_data):
            inbound_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                inbound_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(inbound_table.rows[row_idx].cells[col_idx], "E6E6E1")
                inbound_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()

    h3 = doc.add_heading("4.3 Outbound: Supplier Communication — IN SCOPE", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    outbound_table = doc.add_table(rows=4, cols=3)
    outbound_table.style = 'Table Grid'

    outbound_data = [
        ("Channel", "Topic", "Consumer"),
        ("EDI", "stark.procurement.outbound.pos.edi", "EDI Gateway / Stark Output"),
        ("Email", "stark.procurement.outbound.pos.email", "Stark Output"),
        ("Portal", "stark.procurement.outbound.pos.portal", "Supplier Portal (future)"),
    ]

    for row_idx, row_data in enumerate(outbound_data):
        for col_idx, cell_data in enumerate(row_data):
            outbound_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                outbound_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(outbound_table.rows[row_idx].cells[col_idx], "E6E6E1")
                outbound_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()
    add_horizontal_line(doc)

    # EXPLICIT EXCLUSIONS
    h2 = doc.add_heading("Explicit Exclusions", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("Will NOT Be Delivered", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    exclusions_table = doc.add_table(rows=11, cols=3)
    exclusions_table.style = 'Table Grid'

    exclusions_data = [
        ("Item", "Reason", "Future Phase"),
        ("Stock Transfer Orders (STO)", "Requires NYCE WMS", "Phase 2"),
        ("3-Way Invoice Match", "Requires NYCE goods receipt", "Phase 2"),
        ("SAP Finance Integration", "Requires SAP Finance live", "Phase 2"),
        ("Real-time SalesApp Sync", "Requires OMI bidirectional", "Phase 2"),
        ("Supplier Portal", "Enhancement, not MVP", "Phase 3"),
        ("AI Email Parsing", "Enhancement, not MVP", "Phase 3"),
        ("Contract Management", "Requires Icertis", "Phase 3"),
        ("Mobile Native App", "Enhancement, not MVP", "Phase 3"),
        ("Multi-language UI", "Enhancement, not MVP", "Phase 3"),
        ("Advanced Analytics/BI", "Different project", "Never (BI team)"),
    ]

    for row_idx, row_data in enumerate(exclusions_data):
        for col_idx, cell_data in enumerate(row_data):
            exclusions_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                exclusions_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(exclusions_table.rows[row_idx].cells[col_idx], "E6E6E1")
                exclusions_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()

    h3 = doc.add_heading("Dependencies We Assume Are Ready", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    deps = [
        ("Relex PR feed", "READY — Automated replenishment"),
        ("ECom PR feed", "READY — Drop-shipment PRs"),
        ("Stark Output (EDI/Email)", "READY — Cannot send POs to suppliers without this"),
        ("SSO/IAM", "READY — No user authentication without this"),
        ("Supplier Master Data", "READY — Via MDM"),
    ]

    for dep, status in deps:
        p = doc.add_paragraph()
        run = p.add_run(f"{dep}: ")
        run.bold = True
        p.add_run(status)
        p.paragraph_format.space_after = Pt(6)

    add_horizontal_line(doc)

    # ACCEPTANCE CRITERIA
    h2 = doc.add_heading("Acceptance Criteria Summary", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    h3 = doc.add_heading("Definition of Done (Per Screen)", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    screen_criteria = [
        "Renders correctly on desktop (1280px+)",
        "Renders correctly on tablet (768px+)",
        "All data loads from backend API",
        "Loading states shown during fetch",
        "Error states handled gracefully",
        "Empty states designed and implemented",
        "Keyboard navigation works",
        "Meets WCAG 2.1 AA (basic)",
        "No console errors",
        "Passes TypeScript strict mode",
    ]

    for item in screen_criteria:
        p = doc.add_paragraph(f"☐ {item}")
        p.paragraph_format.space_after = Pt(2)

    h3 = doc.add_heading("Definition of Done (Per Integration)", level=3)
    h3.runs[0].font.color.rgb = AA_DARK_GRAY

    int_criteria = [
        "API endpoint documented (OpenAPI)",
        "Request validation implemented",
        "Error responses follow standard format",
        "Rate limiting configured",
        "Authentication required",
        "Audit logging enabled",
        "Health check includes dependency",
        "Retry logic for transient failures",
    ]

    for item in int_criteria:
        p = doc.add_paragraph(f"☐ {item}")
        p.paragraph_format.space_after = Pt(2)

    add_horizontal_line(doc)

    # SIGN-OFF
    h2 = doc.add_heading("Sign-Off Checklist", level=2)
    h2.runs[0].font.color.rgb = AA_BLACK

    signoff_intro = doc.add_paragraph()
    signoff_intro.add_run("Before final delivery, both parties confirm:")
    signoff_intro.paragraph_format.space_after = Pt(12)

    signoff_table = doc.add_table(rows=9, cols=3)
    signoff_table.style = 'Table Grid'

    signoff_data = [
        ("Item", "Vendor", "STARK"),
        ("All IN SCOPE items delivered", "☐", "☐"),
        ("All BOUNDARY interfaces documented", "☐", "☐"),
        ("All OUT OF SCOPE items acknowledged", "☐", "☐"),
        ("All exclusions agreed", "☐", "☐"),
        ("Acceptance criteria met", "☐", "☐"),
        ("Source code transferred", "☐", "☐"),
        ("Documentation complete", "☐", "☐"),
        ("Training delivered", "☐", "☐"),
    ]

    for row_idx, row_data in enumerate(signoff_data):
        for col_idx, cell_data in enumerate(row_data):
            signoff_table.rows[row_idx].cells[col_idx].text = cell_data
            if row_idx == 0:
                signoff_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                set_cell_shading(signoff_table.rows[row_idx].cells[col_idx], "E6E6E1")
                signoff_table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.color.rgb = AA_BLACK

    doc.add_paragraph()
    doc.add_paragraph()

    # CLOSING
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.add_run("— End of Scope Definition —")
    closing.runs[0].italic = True
    closing.runs[0].font.color.rgb = AA_MID_GRAY

    doc.add_paragraph()

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.add_run("Prepared by Agentic Agency")
    prepared.runs[0].font.size = Pt(10)
    prepared.runs[0].font.color.rgb = AA_MID_GRAY

    tagline_end = doc.add_paragraph()
    tagline_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_end.add_run("Engineering > Prompting")
    tagline_end.runs[0].italic = True
    tagline_end.runs[0].font.size = Pt(10)
    tagline_end.runs[0].font.color.rgb = AA_MID_GRAY

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_document()
